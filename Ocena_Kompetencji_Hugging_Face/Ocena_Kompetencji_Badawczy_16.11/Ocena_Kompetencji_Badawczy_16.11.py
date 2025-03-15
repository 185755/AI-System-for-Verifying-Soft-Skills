## OCENA KOMPETENCJI
from __future__ import print_function
import os.path
import io

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
import requests
import whisper
from huggingface_hub import login
from docx import Document


# Podaj klucz API
api_key = "?????????????"  # Twój token API
# Zalogowanie się przy użyciu tokena
login(api_key)

print("Zalogowano do Hugging Face!")


# Ścieżka do pliku z historią kandydata
file_path = "zadanie1_transkrypcja.txt"



# Wczytaj zawartość pliku
with open(file_path, "r", encoding="utf-8") as file:
    candidate_story = file.read()

# Adres API Hugging Face # https://api-inference.huggingface.co/models/ {tutaj wkleić endpoint}


url = f"https://api-inference.huggingface.co/models/google/flan-t5-base"


# Zdefiniuj funkcję oceny kandydata
def evaluate_candidate_hf(story, api_key):
    # Tworzenie promptu dla modelu w oparciu o metodę STAR
    prompt = (
        f"Zadanie: Opisz sytuację, w której podczas realizacji zadania zawodowego napotkałeś na trudności. "
        f"Co było Twoim zadaniem i jakie trudności napotkałeś?\n"
        f"Jakie działania podjąłeś w obliczu trudności, aby je przezwyciężyć?\n"
        f"Jakie były rezultaty Twoich działań? Oceń całą sytuację w skali od 1 do 10.\n\n"
        f"Historia kandydata:\n{story}\n\nOcena modelu:"
    )

    # Przygotowanie nagłówków dla autoryzacji
    headers = {
        "Authorization": f"Bearer {api_key}"
    }

    # Przygotowanie danych do wysłania w zapytaniu
    payload = {
        "inputs": prompt,
    }

    # Wysłanie zapytania do API Hugging Face
    response = requests.post(url, headers=headers, json=payload, timeout=100)

    if response.status_code == 200:
        return response.json()[0]["generated_text"]  # Zwróć wygenerowany tekst
    else:
        print(f"Błąd: {response.status_code}, {response.text}")
        return None




# Wywołanie oceny
evaluation = evaluate_candidate_hf(candidate_story, api_key)

# Sprawdzenie, czy odpowiedź została poprawnie wygenerowana
if evaluation:
    print("Ocena kandydata:")
    print(evaluation)

    # Wynik w pliku doc
    output_file = "ocena_kandydata.docx"
    document = Document()

    # Dodanie zawartości do pliku
    document.add_heading("Ocena Kandydata", level=1)
    document.add_paragraph("Historia Kandydata:")
    document.add_paragraph(candidate_story)
    document.add_paragraph("\nOcena Modelu:")
    document.add_paragraph(evaluation)

    # Zapisz dokument
    document.save(output_file)

    print(f"Ocena została zapisana w pliku: {output_file}")
else:
    print("Nie udało się wygenerować oceny.")
